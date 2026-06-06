from pydantic import Json
from service.base_service import BaseService
from langchain_core.documents import Document
from pathlib import Path
import os
import subprocess
import platform
from typing import List, Any

# 添加MIME处理所需的导入
import email
from email import policy
from email.parser import BytesParser
import quopri
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
# 改为
from langchain_core.documents import Document
from pathlib import Path
from docx import Document as DocxDocument
from langchain_community.document_loaders import (
    UnstructuredFileLoader,
    PyPDFLoader,
    Docx2txtLoader,
    TextLoader,
    UnstructuredPowerPointLoader,
    UnstructuredExcelLoader,
    CSVLoader
)
from json import dumps
 

class ReadFileService(BaseService):
    def __init__(self):
        super().__init__()
        
    def detect_file_type(self, file_path):
        """检测文件的实际内容类型"""
        try:
            with open(file_path, 'rb') as f:
                # 读取前 1024 字节
                header = f.read(1024)
            
            # 检测 MIME 格式
            mime_signatures = [b'Message-ID:', b'MIME-Version:', b'Content-Type:', b'From:', b'To:']
            if any(sig in header for sig in mime_signatures):
                return 'mime'
            
            # 检测 HTML 文件头
            html_signatures = [b'<!DOCTYPE html', b'<html', b'<HTML', b'<head', b'<HEAD']
            if any(header.startswith(sig) for sig in html_signatures) or any(sig in header for sig in html_signatures):
                return 'html'
            
            # 检测 DOC 文件头（多种情况）
            # 1. 标准 OLE2 Compound Document Format (.doc, .xls, .ppt 等)
            if header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                return 'doc'
            
            # 2. 检查文件扩展名作为辅助判断
            file_extension = os.path.splitext(file_path)[1].lower()
            if file_extension == '.doc':
                # 对于扩展名为.doc的文件，进一步检查内容特征
                # 检查是否包含典型的Word文档内容特征
                doc_signatures = [
                    b'Microsoft Word',
                    b'Word.Document',
                    b'application/msword',
                    b'PK\x03\x04'  # ZIP格式（可能是.docx伪装成.doc）
                ]
                if any(sig in header for sig in doc_signatures):
                    return 'doc'
                
                # 对于无法通过特征检测的.doc文件，直接返回'doc'
                # 因为用户明确指定这是一个.doc文件
                return 'doc'
            
            return 'unknown'
        except Exception as e:
            self.log_error(f"检测文件类型失败: {e}")
            # 如果检测失败，根据扩展名返回默认类型
            file_extension = os.path.splitext(file_path)[1].lower()
            if file_extension == '.doc':
                return 'doc'
            return 'unknown'
    
    def read_html_doc(self, file_path):
        """读取实际是 HTML 内容的 .doc 文件"""
        try:
            # 1. 尝试使用不同编码读取文件
            encodings = ['utf-8', 'gbk', 'gb2312', 'iso-8859-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                raise Exception("无法解码文件内容")
            
            # 2. 使用 BeautifulSoup 提取纯文本
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            
            # 移除脚本和样式
            for script in soup(['script', 'style']):
                script.decompose()
            
            # 获取纯文本
            text = soup.get_text()
            
            # 清理文本
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
        except Exception as e:
            self.log_error(f"读取 HTML 文档失败: {e}")
            raise
    
    def read_mime_doc(self, file_path):
        """读取 MIME 格式的 .doc 文件，提取实际内容"""
        try:
            with open(file_path, 'rb') as f:
                # 解析 MIME 结构
                msg = BytesParser(policy=policy.default).parse(f)
            
            # 提取所有文本内容
            text_parts = []
            
            def extract_part(part):
                """递归提取邮件部分内容"""
                if part.is_multipart():
                    for subpart in part.iter_parts():
                        extract_part(subpart)
                else:
                    content_type = part.get_content_type()
                    charset = part.get_content_charset() or 'utf-8'
                    
                    if content_type in ['text/plain', 'text/html']:
                        # 获取原始内容
                        raw_content = part.get_content()
                        
                        # 如果是 quoted-printable 编码，进行解码
                        if part['Content-Transfer-Encoding'] == 'quoted-printable':
                            # 处理 quoted-printable 编码
                            if isinstance(raw_content, str):
                                # 如果是字符串，先编码为 bytes
                                raw_content = quopri.decodestring(raw_content.encode('utf-8')).decode(charset)
                            else:
                                # 如果已经是 bytes
                                raw_content = quopri.decodestring(raw_content).decode(charset)
                        
                        # 如果是 HTML，提取纯文本
                        if content_type == 'text/html':
                            from bs4 import BeautifulSoup
                            soup = BeautifulSoup(raw_content, 'html.parser')
                            text = soup.get_text()
                        else:
                            text = raw_content
                        
                        if text.strip():
                            text_parts.append(text)
            
            # 开始提取
            extract_part(msg)
            
            if text_parts:
                return '\n'.join(text_parts)
            else:
                raise Exception("未找到文本内容")
        except Exception as e:
            self.log_error(f"读取 MIME 文档失败: {e}")
            raise
    
    def read_doc_via_libreoffice(self, file_path):
        """使用 LibreOffice 将 .doc 转换为 .docx 并读取"""
        # 1. 使用 LibreOffice 命令行将 .doc 转为 .docx
        # --headless 也就是无界面模式
        output_dir = os.path.dirname(file_path)
        subprocess.run([
            'soffice', '--headless', '--convert-to', 'docx', 
            '--outdir', output_dir, file_path
        ], check=True)
        
        # 2. 读取生成的 .docx 文件
        docx_path = file_path + "x" # 假设文件名只是加了x
        from docx import Document
        doc = Document(docx_path)
        
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # 清理临时文件 (可选)
        # os.remove(docx_path)
        
        return '\n'.join(full_text)
    
    def read_doc_with_fallback(self, file_path):
        """使用多种策略读取 .doc 文件"""
        # 1. 首先检测文件实际类型
        file_type = self.detect_file_type(file_path)
        content = None
        if file_type == 'mime':
            self.log_info(f"文件 {file_path} 实际是 MIME 格式，使用 MIME 解析器读取")
            content = self.read_mime_doc(file_path)
        
        if file_type == 'html':
            self.log_info(f"文件 {file_path} 实际是 HTML 格式，使用 HTML 解析器读取")
            # 尝试使用 html2text 读取
            try:
               content = self.read_html_with_html2text(file_path)
            except Exception as e:
                self.log_warning(f"html2text 读取失败: {e}，尝试使用 BeautifulSoup")
                content = self.read_html_doc(file_path)
        #判断是否为空
        if content:
            return [Document(page_content=content, metadata={"source": file_path})]
        else:
            return self.read_docx_file(file_path)
            # # 2. 尝试使用 mammoth 库直接读取（如果可用）
            # try:
            #     with open(file_path, 'rb') as f:
            #         result = mammoth.extract_raw_text(f)
            #         text = result.value
            #         if text.strip():
            #             self.log_info(f"使用 mammoth 库成功读取 .doc 文件")
            #             return text
            # except Exception as e:
            #     self.log_warning(f"mammoth 库读取失败: {e}")
            
            # # 3. 尝试使用 python-docx 直接读取（如果是 docx 伪装的 doc）
            # try:
            #     from docx import Document
            #     doc = Document(file_path)
            #     paragraphs = [p.text for p in doc.paragraphs]
            #     text = '\n'.join(paragraphs)
            #     if text.strip():
            #         self.log_info(f"使用 python-docx 成功读取 .doc 文件")
            #         return text
            # except Exception as e:
            #     self.log_warning(f"python-docx 读取失败: {e}")
            
            # 4. 尝试使用 LibreOffice 转换（当前实现）
        #     try:
        #         # 构建命令 - subprocess会自动处理空格，不需要手动添加引号
        #         cmd = [
        #             libreoffice_cmd,
        #             '--headless',
        #             '--convert-to', 'pdf',
        #             '--outdir', output_dir,
        #             file_path
        #         ]
                
        #         self.log_info(f"执行Word转换命令: {' '.join(cmd)}")
        #         result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)  # 
        #         # text = self.read_doc_via_libreoffice(file_path)
        #         if text.strip():
        #             self.log_info(f"使用 LibreOffice 转换成功读取 .doc 文件")
        #             return text
        #     except Exception as e:
        #         self.log_warning(f"LibreOffice 转换失败: {e}")
            
        #     # 5. 尝试使用 UnstructuredFileLoader
        #     try:
        #         from langchain_community.document_loaders import UnstructuredFileLoader
        #         loader = UnstructuredFileLoader(file_path, strategy="fast")
        #         docs = loader.load()
        #         if docs and docs[0].page_content.strip():
        #             self.log_info(f"使用 UnstructuredFileLoader 成功读取 .doc 文件")
        #             return docs[0].page_content
        #     except Exception as e:
        #         self.log_warning(f"UnstructuredFileLoader 读取失败: {e}")
            
        #     # 6. 最终兜底：尝试直接读取为文本
        #     try:
        #         with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        #             text = f.read()
        #         if text.strip():
        #             self.log_info(f"使用直接文本读取成功读取 .doc 文件")
        #             return text
        #     except Exception as e:
        #         self.log_warning(f"直接文本读取失败: {e}")
            
        #     raise Exception("所有读取策略都失败")
        # except Exception as e:
        #     self.log_error(f"读取 .doc 文件失败: {e}")
        #     raise
    
    def get_libreoffice_cmd(self):
        """根据操作系统获取 LibreOffice 命令"""
        if platform.system() == "Windows":
            return r"C:\Program Files\LibreOffice\program\soffice.exe"
        else:
            return "soffice"
    
    def read_pdf_file(self, file_path: str) -> List[Document]:
        """读取 PDF 文件"""
        from langchain_community.document_loaders import PyPDFLoader
        self.log_info(f"使用 PyPDFLoader 加载 PDF 文件")
        loader = PyPDFLoader(file_path)
        return loader.load()
    
    def read_txt_file(self, file_path: str) -> List[Document]:
        """读取 TXT 文件"""
        from langchain_community.document_loaders import TextLoader
        self.log_info(f"使用 TextLoader 加载 TXT 文件")
        
        # 尝试不同的编码，优先使用 UTF-8
        try:
            # 优先使用 UTF-8 编码
            loader = TextLoader(file_path, encoding='utf-8')
            return loader.load()
        except UnicodeDecodeError:
            self.log_warning(f"UTF-8 编码读取失败，尝试使用 GBK 编码")
            # 如果 UTF-8 失败，尝试使用 GBK
            loader = TextLoader(file_path, encoding='gbk')
            return loader.load()
        except Exception as e:
            self.log_warning(f"TextLoader 加载失败 ({e})，使用 UnstructuredFileLoader")
            # 如果都失败，使用 UnstructuredFileLoader
            from langchain_community.document_loaders import UnstructuredFileLoader
            loader = UnstructuredFileLoader(file_path, strategy="fast")
            return loader.load()
    
    def read_doc_file(self, file_path: str) -> List[Document]:
        """读取 DOC 文件"""
        return self.read_doc_with_fallback(file_path)
        
    
    def read_docx_file(self, file_path: str) -> List[Document]:
        """读取 DOCX 文件"""
        loader = None
        libreoffice_cmd = self.get_libreoffice_cmd()
        
        try:
            # 1. 检查源文件是否存在
            if not os.path.exists(file_path):
                raise Exception(f"源文件不存在: {file_path}")
            if not os.access(file_path, os.R_OK):
                raise Exception(f"无读取权限: {file_path}")
            
            # 2. 生成PDF路径
            pdf_path = Path(file_path).with_suffix('.pdf')
            self.log_info(f"检测到 DOCX 文档，尝试转换为 PDF...")
            
            # 3. 优化转换命令
            output_dir = str(Path(file_path).parent)
            
            # 检查输出目录权限
            if not os.access(output_dir, os.W_OK):
                raise Exception(f"无写入权限: {output_dir}")
            
            # 构建命令 - subprocess会自动处理空格，不需要手动添加引号
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--nologo',
                '--norestore',
                '--nolockcheck',
                '--convert-to', 'pdf:writer_pdf_Export',
                '--outdir', output_dir,
                file_path
            ]
            
            self.log_info(f"执行Word转换命令: {' '.join(cmd)}")
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)  # 延长超时时间
            
            self.log_info(f"Word转换命令返回码: {result.returncode}")
            self.log_info(f"Word转换命令stdout: {result.stdout}")
            self.log_info(f"Word转换命令stderr: {result.stderr}")
            
            # 4. 只信任实际生成的文件，忽略返回码
            if pdf_path.exists() and pdf_path.stat().st_size > 0:
                self.log_info(f"✓ word转换pdf成功，使用 PyPDFLoader 加载: {pdf_path}")
                loader = PyPDFLoader(str(pdf_path))
            else:
                # 无论返回码如何，只要PDF未生成就是失败
                result_info = f"Word转换失败 - 源文件: {file_path}, 返回码: {result.returncode}, 错误: {result.stderr}"
                self.log_error(result_info)
                raise Exception(result_info)
                
        except Exception as e:
            self.log_warning(f"PDF 转换失败 ({e})，回退到其他加载器（无页码信息）")
            try:
                loader = Docx2txtLoader(file_path)
                self.log_info(f"使用 Docx2txtLoader 加载 Word 文件（注意：无页码信息）")
            except Exception as e2:
                self.log_warning(f"加载器也失败 ({e2})，使用 UnstructuredFileLoader")
                loader = UnstructuredFileLoader(file_path, strategy="fast")
        documents = loader.load()
        return documents
    
    def read_ppt_file(self, file_path: str) -> List[Document]:
        """读取 PPT/PPTX 文件"""
        documents = []
        libreoffice_cmd = self.get_libreoffice_cmd()
        
        try:
            # 1. 检查源文件是否存在
            if not os.path.exists(file_path):
                raise Exception(f"源文件不存在: {file_path}")
            if not os.access(file_path, os.R_OK):
                raise Exception(f"无读取权限: {file_path}")
            
            # 2. 生成PDF路径
            pdf_path = Path(file_path).with_suffix('.pdf')
            self.log_info(f"检测到 PPT 文档，尝试转换为 PDF...")
            
            # 3. 优化转换命令
            output_dir = str(Path(file_path).parent)
            
            # 检查输出目录权限
            if not os.access(output_dir, os.W_OK):
                raise Exception(f"无写入权限: {output_dir}")
            
            # 构建命令 - subprocess会自动处理空格，不需要手动添加引号
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--nologo',
                '--norestore',
                '--nolockcheck',
                '--convert-to', 'pdf:writer_pdf_Export',
                '--infilter=Impress 2007',
                '--outdir', output_dir,
                file_path
            ]
            
            self.log_info(f"执行转换命令: {' '.join(cmd)}")
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)  # 延长超时时间
            
            self.log_info(f"转换命令返回码: {result.returncode}")
            self.log_info(f"转换命令stdout: {result.stdout}")
            self.log_info(f"转换命令stderr: {result.stderr}")
            
            # 4. 只信任实际生成的文件，忽略返回码
            if pdf_path.exists() and pdf_path.stat().st_size > 0:
                self.log_info("✓ PPT转换PDF成功，使用 PyPDFLoader 加载")
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(str(pdf_path))
                return loader.load()
            else:
                # 无论返回码如何，只要PDF未生成就是失败
                result_info = f"PPT转换失败 - 源文件: {file_path}, 返回码: {result.returncode}, 错误: {result.stderr}"
                self.log_error(result_info)
                raise Exception(result_info)
        except Exception as e:
            self.log_warning(f"PPT转换PDF失败 ({e})，回退到 UnstructuredPowerPointLoader")
            try:
                loader = UnstructuredPowerPointLoader(file_path)
                self.log_info(f"使用 UnstructuredPowerPointLoader 加载 PPT 文件（注意：无页码信息）")
                return loader.load()
            except Exception as e2:
                self.log_warning(f"加载器也失败 ({e2})，使用 UnstructuredFileLoader")
                from langchain_community.document_loaders import UnstructuredFileLoader
                loader = UnstructuredFileLoader(file_path, strategy="fast")
                return loader.load()
    
    def read_csv_file(self, file_path: str) -> List[Document]:
        """读取 CSV 文件"""
        try:
            from langchain_community.document_loaders import CSVLoader
            loader = CSVLoader(file_path, encoding='utf-8')
            self.log_info(f"使用 CSVLoader 加载 CSV 文件")
            return loader.load()
        except UnicodeDecodeError:
            self.log_warning("UTF-8 解码失败，尝试 GBK")
            from langchain_community.document_loaders import CSVLoader
            loader = CSVLoader(file_path, encoding='gbk')
            return loader.load()
        except Exception as e:
            self.log_warning(f"CSVLoader 加载失败 ({e})，使用 UnstructuredFileLoader")
            from langchain_community.document_loaders import UnstructuredFileLoader
            loader = UnstructuredFileLoader(file_path, strategy="fast")
            return loader.load()
    
    def read_excel_file(self, file_path: str) -> List[Document]:
        """读取 Excel 文件"""
        try:
            loader = UnstructuredExcelLoader(file_path)
            self.log_info(f"使用 UnstructuredExcelLoader 加载 Excel 文件（推荐用于embedding）")
            #打印loader.load()
            self.log_info(f"loader.load(): {loader.load()}")
            # self.log_info(dumps(loader.load(), ensure_ascii=False, indent=2))
            return loader.load()
        except Exception as e:
            self.log_error(f"UnstructuredExcelLoader处理失败: {e}")
            loader = UnstructuredFileLoader(file_path, strategy="fast")
            self.log_info(f"使用 UnstructuredFileLoader 加载 Excel 文件（推荐用于embedding）")
            #打印loader.load()
            self.log_info(f"loader.load(): {loader.load()}")
            # self.log_info(dumps(loader.load(), ensure_ascii=False, indent=2))
            return loader.load()
    
    def read_file(self, file_path: str) -> List[Document]:
        """根据文件扩展名选择合适的加载器读取文件"""
        file_extension = os.path.splitext(file_path)[1].lower()
        self.log_info(f"准备加载文件: {file_path}, 扩展名: {file_extension}")
        
        if file_extension == '.pdf':
            return self.read_pdf_file(file_path)
        elif file_extension in ['.txt']:
            return self.read_txt_file(file_path)
        elif file_extension in ['.doc']:
            return self.read_doc_file(file_path)
        elif file_extension in ['.docx']:
            return self.read_docx_file(file_path)
        elif file_extension in ['.pptx', '.ppt']:
            return self.read_ppt_file(file_path)
        elif file_extension == '.csv':
            return self.read_csv_file(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            return self.read_excel_file(file_path)
        else:
            self.log_info(f"当前文件类型不支持，{file_extension}")
            return []
