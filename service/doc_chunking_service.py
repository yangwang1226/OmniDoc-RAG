from xxlimited import Str
from service.base_service import BaseService
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
# 改为
from langchain_core.documents import Document
from pathlib import Path
from docx import Document as DocxDocument
from langchain_community.document_loaders import (
    UnstructuredFileLoader,
    PyPDFLoader,
    Docx2txtLoader,
    TextLoader,
    UnstructuredPowerPointLoader,
    UnstructuredExcelLoader,
    CSVLoader
)
import os
import re
from typing import Dict, Any
from service.read_file_service import ReadFileService


class DocChunkingService(BaseService):
    def __init__(self):
        super().__init__()
        
        # 初始化文件读取服务
        self.read_file_service = ReadFileService()
        
        # 父块分割器（大块，保留完整上下文）
        self.parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,  # 更大的块
            chunk_overlap=200,
            # 增加兜底分隔符，防止因找不到Markdown标题导致无法切分
            separators=["\n## ", "\n### ", "\n\n", "\n", "。", "；", "，", " ", ""],  # 增加更多分隔符
            keep_separator=True,
            is_separator_regex=False
        )
        
        # 子块分割器（小块，用于精确检索）
        self.child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=400,   # 更小的块
            chunk_overlap=50,
            separators=["\n## ", "\n### ", "\n\n", "\n", "。", "；", "，", " ", ""],  # 增加更多分隔符
            keep_separator=True,
            is_separator_regex=False
        )

    # 分块内容
    def chunk_content(self, content: str) -> Dict[str, Any]:
        """分块内容"""
        try:
            parent_chunks = self.parent_splitter.split_text(content)

            # 3. 为每个父块生成子块
            all_child_chunks = []
            
            for parent_chunk in parent_chunks:
                # 从父块中切分子块
                child_chunks = self.child_splitter.split_text(parent_chunk)
                # 为每个子块添加父块引用
                for child_text in child_chunks:
                    # 确保 parent_content 不超过数据库限制
                    safe_parent_content = parent_chunk[:8000] if len(parent_chunk) > 8000 else parent_chunk
                    
                    child_doc = Document(
                        page_content=child_text,
                        metadata={
                            "parent_content": safe_parent_content,
                        }
                    )
                    all_child_chunks.append(child_doc)
            
            self.log_info(f"生成 {len(all_child_chunks)} 个子块")
            
            return {
                "parent_chunks": parent_chunks,
                "child_chunks": all_child_chunks
            }
        except Exception as e:
            self.log_error("分块内容失败", e)
            raise Exception(f"分块内容失败: {str(e)}")

    # def read_doc_via_libreoffice(self, file_path):
    #     # 1. 使用 LibreOffice 命令行将 .doc 转为 .docx
    #     # --headless 也就是无界面模式
    #     output_dir = os.path.dirname(file_path)
    #     import subprocess
    #     subprocess.run([
    #         'soffice', '--headless', '--convert-to', 'docx', 
    #         '--outdir', output_dir, file_path
    #     ], check=True)
        
    #     # 2. 读取生成的 .docx 文件
    #     docx_path = file_path + "x" # 假设文件名只是加了x
    #     doc = docx.Document(docx_path)
        
    #     full_text = []
    #     for para in doc.paragraphs:
    #         full_text.append(para.text)
            
    #     # 清理临时文件 (可选)
    #     # os.remove(docx_path)
        
    #     return '\n'.join(full_text)
    # 使用父子块策略分块文档
    def chunk_document_with_hierarchy(self, actual_file_path: str):
        """使用父子块策略分块文档"""
        documents = []
        try:
            # 1. 使用统一的文件读取服务加载文件
            documents = self.read_file_service.read_file(actual_file_path)
            
            self.log_info(f"加载文档成功，共 {len(documents)} 页")
            
            if not documents or not documents[0].page_content.strip():
                raise ValueError("文档内容为空")
            
            # 判断是否为影印版本PDF（扫描件）
            file_extension = os.path.splitext(actual_file_path)[1].lower()
            if file_extension == '.pdf':
                total_pages = len(documents)
                total_text_length = sum(len(doc.page_content.strip()) for doc in documents)
                avg_text_per_page = total_text_length / total_pages if total_pages > 0 else 0
                
                # 计算非空白字符比例
                non_whitespace_chars = sum(len(re.sub(r'\s+', '', doc.page_content)) for doc in documents)
                non_whitespace_ratio = non_whitespace_chars / total_text_length if total_text_length > 0 else 0
                
                # 影印版特征：平均每页文本少，或非空白字符比例低
                is_scanned = avg_text_per_page < 100 or (avg_text_per_page < 300 and non_whitespace_ratio < 0.5)
                
                if is_scanned:
                    self.log_warning(f"检测到可能是影印版本PDF，平均每页文本长度: {avg_text_per_page:.2f}，非空白字符比例: {non_whitespace_ratio:.2f}")
                    raise ValueError("检测到可能是影印版本PDF，不支持处理")
                    # 可以在这里添加处理逻辑，比如调用OCR服务
            
            # 2. 先切分父块
            parent_chunks = self.parent_splitter.split_documents(documents)
            self.log_info(f"生成 {len(parent_chunks)} 个父块")
            
            # 3. 为每个父块生成子块
            all_child_chunks = []
            
            for parent_idx, parent_chunk in enumerate(parent_chunks):
                # 从父块中切分子块
                child_chunks = self.child_splitter.split_text(parent_chunk.page_content)
                # 获取页码，如果没有则使用0
                page_number = parent_chunk.metadata.get("page")
                page_number = page_number + 1 if page_number is not None else 0
                # 为每个子块添加父块引用
                for child_idx, child_text in enumerate(child_chunks):
                    # 确保 parent_content 不超过数据库限制
                    p_content = parent_chunk.page_content
                    safe_parent_content = p_content[:8000] if len(p_content) > 8000 else p_content
                    
                    child_doc = Document(
                        page_content=child_text,
                        metadata={
                            "page_number": page_number,
                            "parent_content": safe_parent_content,
                        }
                    )
                    all_child_chunks.append(child_doc)
            
            self.log_info(f"生成 {len(all_child_chunks)} 个子块")
            
            return {
                "parent_chunks": parent_chunks,
                "child_chunks": all_child_chunks
            }
            
        except Exception as e:
            self.log_error("父子块处理失败", e)
            raise
        finally:
            # 清空documents列表，释放内存
            documents.clear()
            self.log_info("资源清理完成")
        # 使用父子块策略分块文档
    # def chunk_document_with_hierarchy(self, actual_file_path: Str):
    #     """使用父子块策略分块文档"""
    #     try:
    #         # 1. 根据文件扩展名选择合适的加载器
    #         file_extension = os.path.splitext(actual_file_path)[1].lower()
    #         self.log_info(f"准备加载文件: {actual_file_path}, 扩展名: {file_extension}")
            
    #         loader = None
    #         import subprocess
    #         import platform
            
    #         # 根据操作系统选择命令
    #         if platform.system() == "Windows":
    #             libreoffice_cmd = r"C:\Program Files\LibreOffice\program\soffice.exe"
    #         else:
    #             libreoffice_cmd = "soffice"
    #         if file_extension == '.pdf':
    #             loader = PyPDFLoader(actual_file_path)
    #             self.log_info(f"使用 PyPDFLoader 加载 PDF 文件")
    #         elif file_extension in ['.docx', '.doc']:
    #             try:
    #                 if not pdf_path.exists():
    #                     output_dir = str(Path(actual_file_path).parent)
                        
    #                     # 检查输出目录权限
    #                     if not os.access(output_dir, os.W_OK):
    #                         raise Exception(f"无写入权限: {output_dir}")
                        
    #                     # 构建命令 - subprocess会自动处理空格，不需要手动添加引号
    #                     cmd = [
    #                         libreoffice_cmd,
    #                         '--headless',
    #                         '--convert-to', 'pdf',
    #                         '--outdir', output_dir,
    #                         actual_file_path
    #                     ]
                        
    #                     self.log_info(f"执行Word转换命令: {' '.join(cmd)}")
    #                     result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)  # 延长超时时间
                        
    #                     self.log_info(f"Word转换命令返回码: {result.returncode}")
    #                     self.log_info(f"Word转换命令stdout: {result.stdout}")
    #                     self.log_info(f"Word转换命令stderr: {result.stderr}")
                        
    #                     # 只信任实际生成的文件，忽略返回码
    #                     if pdf_path.exists() and pdf_path.stat().st_size > 0:
    #                         self.log_info("✓ word转换pdf成功，使用 PyPDFLoader 加载")
    #                         loader = PyPDFLoader(str(pdf_path))
    #                     else:
    #                         # 无论返回码如何，只要PDF未生成就是失败
    #                         result_info = f"Word转换失败 - 源文件: {actual_file_path}, 返回码: {result.returncode}, 错误: {result.stderr}"
    #                         self.log_error(result_info)
    #                         raise Exception(result_info)
                        
    #             except Exception as e:
    #                 self.log_warning(f"PDF 转换失败 ({e})，回退到其他加载器（无页码信息）")
    #                 try:
    #                     loader = Docx2txtLoader(actual_file_path)
    #                     self.log_info(f"使用 Docx2txtLoader 加载 Word 文件（注意：无页码信息）")
    #                 except Exception as e2:
    #                     self.log_warning(f"加载器也失败 ({e2})，使用 UnstructuredFileLoader")
    #                     loader = UnstructuredFileLoader(actual_file_path, strategy="fast")
    #         elif file_extension in ['.pptx', '.ppt']:
    #             try:
    #                 # 1. 检查源文件是否存在
    #                 if not os.path.exists(actual_file_path):
    #                     raise Exception(f"源文件不存在: {actual_file_path}")
    #                 if not os.access(actual_file_path, os.R_OK):
    #                     raise Exception(f"无读取权限: {actual_file_path}")
                    
    #                 # 2. 生成PDF路径
    #                 pdf_path = Path(actual_file_path).with_suffix('.pdf')
    #                 self.log_info(f"检测到 PPT 文档，尝试转换为 PDF...")
                    
    #                 # 3. 优化转换命令
    #                 output_dir = str(Path(actual_file_path).parent)
                    
    #                 # 检查输出目录权限
    #                 if not os.access(output_dir, os.W_OK):
    #                     raise Exception(f"无写入权限: {output_dir}")
                    
    #                 # 构建命令 - subprocess会自动处理空格，不需要手动添加引号
    #                 cmd = [
    #                     libreoffice_cmd,
    #                     '--headless',
    #                     '--nologo',
    #                     '--norestore',
    #                     '--nolockcheck',
    #                     '--convert-to', 'pdf:writer_pdf_Export',
    #                     '--infilter=Impress 2007',
    #                     '--outdir', output_dir,
    #                     actual_file_path
    #                 ]
                    
    #                 self.log_info(f"执行转换命令: {' '.join(cmd)}")
    #                 result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)  # 延长超时时间
                    
    #                 self.log_info(f"转换命令返回码: {result.returncode}")
    #                 self.log_info(f"转换命令stdout: {result.stdout}")
    #                 self.log_info(f"转换命令stderr: {result.stderr}")
                    
    #                 # 4. 只信任实际生成的文件，忽略返回码
    #                 if pdf_path.exists() and pdf_path.stat().st_size > 0:
    #                     self.log_info("✓ PPT转换PDF成功，使用 PyPDFLoader 加载")
    #                     loader = PyPDFLoader(str(pdf_path))
    #                 else:
    #                     # 无论返回码如何，只要PDF未生成就是失败
    #                     result_info = f"PPT转换失败 - 源文件: {actual_file_path}, 返回码: {result.returncode}, 错误: {result.stderr}"
    #                     self.log_error(result_info)
    #                     raise Exception(result_info)
    #             except Exception as e:
    #                 self.log_warning(f"PPT转换PDF失败 ({e})，回退到 UnstructuredPowerPointLoader")
    #                 try:
    #                     loader = UnstructuredPowerPointLoader(actual_file_path)
    #                     self.log_info(f"使用 UnstructuredPowerPointLoader 加载 PPT 文件（注意：无页码信息）")
    #                 except Exception as e2:
    #                     self.log_warning(f"加载器也失败 ({e2})，使用 UnstructuredFileLoader")
    #                     loader = UnstructuredFileLoader(actual_file_path, strategy="fast")

    #         elif file_extension in ['.xlsx', '.xls']:
    #             # 优先使用UnstructuredExcelLoader（适合embedding场景，保持表格结构）
    #             try:
    #                 loader = UnstructuredExcelLoader(actual_file_path)
    #                 self.log_info(f"使用 UnstructuredExcelLoader 加载 Excel 文件（推荐用于embedding）")       
    #             except Exception as e:
    #                 self.log_error(f"Excel处理失败: {e}")
    #                 raise
    #         elif file_extension == '.csv':
    #             try:
    #                 loader = CSVLoader(actual_file_path, encoding='utf-8')
    #                 self.log_info(f"使用 CSVLoader 加载 CSV 文件")
    #             except UnicodeDecodeError:
    #                 self.log_warning("UTF-8 解码失败，尝试 GBK")
    #                 loader = CSVLoader(actual_file_path, encoding='gbk')
    #             except Exception as e:
    #                 self.log_warning(f"CSVLoader 加载失败 ({e})，使用 UnstructuredFileLoader")
    #                 loader = UnstructuredFileLoader(actual_file_path, strategy="fast")
    #         elif file_extension in ['.txt', '.md']:
    #             try:
    #                 loader = TextLoader(actual_file_path, encoding='utf-8')
    #             except UnicodeDecodeError:
    #                 self.log_warning("UTF-8 解码失败，尝试 GBK")
    #                 loader = TextLoader(actual_file_path, encoding='gbk')
    #             self.log_info(f"使用 TextLoader 加载文本文件")
    #         else:
    #             self.log_info(f"当前文件类型不支持，{file_extension}")
    #             # # 其他文件类型使用快速策略
                             
    #         documents = loader.load()
    #         self.log_info(f"加载文档成功，共 {len(documents)} 页")
            
    #         if not documents or not documents[0].page_content.strip():
    #             raise ValueError("文档内容为空")
            
    #         # 判断是否为影印版本PDF（扫描件）
    #         if file_extension == '.pdf':
    #             total_pages = len(documents)
    #             total_text_length = sum(len(doc.page_content.strip()) for doc in documents)
    #             avg_text_per_page = total_text_length / total_pages if total_pages > 0 else 0
                
    #             # 计算非空白字符比例
    #             non_whitespace_chars = sum(len(re.sub(r'\s+', '', doc.page_content)) for doc in documents)
    #             non_whitespace_ratio = non_whitespace_chars / total_text_length if total_text_length > 0 else 0
                
    #             # 影印版特征：平均每页文本少，或非空白字符比例低
    #             is_scanned = avg_text_per_page < 100 or (avg_text_per_page < 300 and non_whitespace_ratio < 0.5)
                
    #             if is_scanned:
    #                 self.log_warning(f"检测到可能是影印版本PDF，平均每页文本长度: {avg_text_per_page:.2f}，非空白字符比例: {non_whitespace_ratio:.2f}")
    #                 raise ValueError("检测到可能是影印版本PDF，不支持处理")
    #                 # 可以在这里添加处理逻辑，比如调用OCR服务
            
    #         # 2. 先切分父块
    #         parent_chunks = self.parent_splitter.split_documents(documents)
    #         self.log_info(f"生成 {len(parent_chunks)} 个父块")
            
    #         # 3. 为每个父块生成子块
    #         all_child_chunks = []
            
    #         for parent_idx, parent_chunk in enumerate(parent_chunks):
    #             # 从父块中切分子块
    #             child_chunks = self.child_splitter.split_text(parent_chunk.page_content)
    #             # 获取页码，如果没有则使用0
    #             page_number = parent_chunk.metadata.get("page")
    #             page_number = page_number + 1 if page_number is not None else 0
    #             # 为每个子块添加父块引用
    #             for child_idx, child_text in enumerate(child_chunks):
    #                 # 确保 parent_content 不超过数据库限制
    #                 p_content = parent_chunk.page_content
    #                 safe_parent_content = p_content[:8000] if len(p_content) > 8000 else p_content
                    
    #                 child_doc = Document(
    #                     page_content=child_text,
    #                     metadata={
    #                         "page_number": page_number,
    #                         "parent_content": safe_parent_content,
    #                     }
    #                 )
    #                 all_child_chunks.append(child_doc)
            
    #         self.log_info(f"生成 {len(all_child_chunks)} 个子块")
            
    #         return {
    #             "parent_chunks": parent_chunks,
    #             "child_chunks": all_child_chunks
    #         }
            
    #     except Exception as e:
    #         self.log_error("父子块处理失败", e)
    #         raise

